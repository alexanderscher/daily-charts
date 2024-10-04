from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException
import time
import pandas as pd
import os
import datetime
import sys
from tempfile import mkdtemp
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

pd.set_option("display.max_rows", None)
pd.set_option("display.max_columns", None)
pd.set_option("display.width", 1000)
pd.set_option("display.colheader_justify", "center")
pd.set_option("display.precision", 3)

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from layers.db.get_db import FetchDB
from layers.utils.check import check_prod
from layers.apis.spotify_api import SpotifyAPI
from layers.utils.check import smart_partial_match

CLIENT_ID = os.getenv("SPOTIFY_CLIENT_ID")
USER_ID = os.getenv("SPOTIFY_USER_ID")
CLIENT_SECRET = os.getenv("SPOTIFY_CLIENT_SECRET")


class Scrape:
    def __init__(self, driver):
        self.db = FetchDB()
        self.roster_artists = self.db.get_roster_artists()
        self.df = []
        self.us = []
        self.pub_songs = self.db.get_pub_songs()
        self.pub_artists = self.db.get_pub_artists()
        self.driver = driver
        self.major_labels = self.db.get_major_labels()
        self.signed_artists = self.db.get_signed_artists()
        self.client = SpotifyAPI(CLIENT_ID, USER_ID, CLIENT_SECRET)
        self.l2tk_chart = []
        self.other = []
        self.prospect_list = []

    def spotify_signin(self):
        try:
            self.driver.get("https://accounts.spotify.com/en/login")
            self.driver.refresh()
            from dotenv import dotenv_values

            config = dotenv_values("/Users/al/Desktop/L2TK.nosync/.env")
            credentials = {
                "username": config["S4A_ALEX_USERNAME"],
                "password": config["S4A_ALEX_PASSWORD"],
            }
            input_username = self.driver.find_element(
                By.XPATH, '//*[@id="login-username"]'
            )
            input_password = self.driver.find_element(
                By.XPATH, '//*[@id="login-password"]'
            )
            time.sleep(2)

            input_username.click()
            input_username.clear()
            input_username.send_keys(credentials["username"])
            time.sleep(2)
            input_password.click()
            input_password.clear()
            input_password.send_keys(credentials["password"])
            time.sleep(2)
            button_log_in = self.driver.find_element(
                By.CLASS_NAME, "ButtonInner-sc-14ud5tc-0"
            )
            button_log_in.click()
            time.sleep(10)
            print("Logged in")
        except NoSuchElementException:
            self.spotify_signin()

    def check_roster(
        self, name, position, artist, track, ch, days, peak, date, callback
    ):
        checked_pub = check_prod(self.pub_songs, self.pub_artists, track, artist)
        artist_exists = any(
            art.lower() in artist.lower() for art in self.roster_artists
        )

        if checked_pub or artist_exists:

            self.df.append((name, position, artist, track, ch, days, peak, date))

        elif callback == None:
            pass
        else:
            callback()

    def spotify(self, name, url):

        self.driver.get(url)
        time.sleep(5)

        n = self.driver.find_element(
            By.XPATH, '//*[@id="__next"]/div/div[3]/div/div/div[1]/h1/span'
        )

        d = self.driver.find_element(By.XPATH, '//*[@id="date_picker"]')
        date = d.get_attribute("value")
        dvg = self.driver.find_elements(By.TAG_NAME, "tr")
        dvg_rows = len(dvg) - 1
        dvg_table_data = self.driver.find_elements(By.TAG_NAME, "td")
        dvg_table_data_length = len(dvg_table_data)
        dvg_columns = int(dvg_table_data_length / dvg_rows)

        for i, song in enumerate(dvg[1 : len(dvg)]):
            position = (
                dvg_table_data[i * dvg_columns + 1]
                .find_elements(By.XPATH, ".//span")[0]
                .text
            )
            peak = dvg_table_data[i * dvg_columns + 3].text
            prev = dvg_table_data[i * dvg_columns + 4].text
            days = dvg_table_data[i * dvg_columns + 5].text
            mov = (
                dvg_table_data[i * dvg_columns + 1]
                .find_elements(By.XPATH, ".//span")[1]
                .text
            )

            try:
                ch = int(prev) - int(position)
                if ch > 0:
                    ch = "+" f"{ch}"
            except ValueError:
                if peak != position and prev == "—":
                    ch = "RE-ENTRY"
                else:
                    ch = "NEW"

            song_anchors = dvg_table_data[i * dvg_columns + 2].find_elements(
                By.XPATH, ".//a"
            )
            track = song_anchors[1].text
            artist = song_anchors[2].text

            def callback():
                if peak == position:
                    mov = (
                        dvg_table_data[i * dvg_columns + 1]
                        .find_elements(By.XPATH, ".//span")[1]
                        .text
                    )
                    no_bueno = " – "
                    if mov != no_bueno:
                        self.df.append(
                            (name, position, artist, track, ch, None, None, date)
                        )

            if peak == position:
                mov = (
                    dvg_table_data[i * dvg_columns + 1]
                    .find_elements(By.XPATH, ".//span")[1]
                    .text
                )
                no_bueno = " – "
                if mov != no_bueno:
                    self.check_roster(
                        name,
                        position,
                        artist,
                        track,
                        f"{ch} (peak)",
                        days,
                        peak,
                        date,
                        callback,
                    )
                else:
                    self.check_roster(
                        name, position, artist, track, f"{ch}", days, peak, date, None
                    )
            else:
                self.check_roster(
                    name, position, artist, track, f"{ch}", days, peak, date, None
                )

    def chart_search(self):
        spotify_df = pd.DataFrame(
            self.df,
            columns=[
                "Chart",
                "Position",
                "Artist",
                "Song",
                "Movement",
                "Days",
                "Peak",
                "Date",
            ],
        )

        for (
            chart,
            position,
            artist,
            song,
            movement,
            days,
            peak,
            date,
        ) in spotify_df.itertuples(index=False):

            if ", " in artist:
                a = artist.split(", ")
                artist = a[0]

            elif " featuring " in artist:
                a = artist.split(" featuring ")
                artist = a[0]

            elif "feat." in artist:
                a = artist.split("feat.")
                artist = a[0]

            checked_pub = check_prod(self.pub_songs, self.pub_artists, song, artist)
            artist_exists = any(
                art.lower() in artist.lower() for art in self.roster_artists
            )

            if checked_pub or artist_exists:
                self.us.append(
                    (
                        chart,
                        position,
                        artist,
                        song,
                        None,
                        "L2TK",
                        movement,
                        days,
                        peak,
                        None,
                        None,
                        date,
                    )
                )

            elif list(
                filter(lambda x: (x.lower() == artist.lower()), self.signed_artists)
            ):
                self.us.append(
                    (
                        chart,
                        position,
                        artist,
                        song,
                        None,
                        None,
                        movement,
                        None,
                        None,
                        None,
                        None,
                        date,
                    )
                )

            else:
                if ", " in artist.lower():
                    a = artist(", ")
                    artist = a[0]
                    copyright = self.client.get_artist_copy_track(
                        artist.lower(), song, "spotify"
                    )

                elif " featuring " in artist.lower():
                    a = artist(" featuring ")
                    artist = a[0]
                    copyright = self.client.get_artist_copy_track(
                        artist.lower(), song, "spotify"
                    )

                elif "feat." in artist.lower():
                    a = artist("feat.")
                    artist = a[0]
                    copyright = self.client.get_artist_copy_track(
                        artist.lower(), song, "spotify"
                    )
                else:

                    copyright = self.client.get_artist_copy_track(
                        artist.lower(), song, "spotify"
                    )

                if copyright:
                    matched_labels = list(
                        filter(
                            lambda x: smart_partial_match(x, copyright[0].lower()),
                            self.major_labels,
                        )
                    )

                    if not matched_labels:

                        self.us.append(
                            (
                                chart,
                                position,
                                artist,
                                song,
                                "UNSIGNED",
                                None,
                                movement,
                                None,
                                None,
                                copyright[1],
                                copyright[0],
                                date,
                            )
                        )

                    else:
                        self.us.append(
                            (
                                chart,
                                position,
                                artist,
                                song,
                                None,
                                None,
                                movement,
                                None,
                                None,
                                None,
                                None,
                                date,
                            )
                        )

    def create_docx(self):
        final_df = unsigned = pd.DataFrame(
            self.us,
            columns=[
                "Chart",
                "Position",
                "Artist",
                "Song",
                "Unsigned",
                "L2TK",
                "Movement",
                "Days",
                "Peak",
                "Link",
                "Label",
                "Date",
            ],
        )
        document = Document()

        document.add_paragraph(
            "Spotify Chart Report - "
            + datetime.datetime.now().strftime("%m/%d/%y")
            + "\n Conor Ambrose <conor@listen2thekids.com>,  Ari Kononov <ari@listen2thekids.com>, laura@listen2thekids.com \n"
        )

        chart_header = None

        def add_content_and_header(chart, date):
            if self.l2tk_chart or self.other or self.prospect_list:
                para_header = document.add_paragraph()
                font = para_header.style.font
                font.name = "Arial"
                font.size = Pt(9)
                header_text = (
                    "\n" + f"{chart}\n".upper()
                    if pd.isna(date)
                    else "\n" + f"{chart} - {date}\n".upper()
                )
                runner = para_header.add_run(header_text)
                runner.bold = True
                runner.underline = True
                if self.l2tk_chart:
                    document.add_paragraph("L2TK:")
                    for p in self.l2tk_chart:
                        para = document.add_paragraph(p)
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

                if self.prospect_list:
                    document.add_paragraph("PROSPECT:")
                    for p in self.prospect_list:
                        para = document.add_paragraph(p)
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.RED

                if self.other:
                    document.add_paragraph("NEW ADDS:")
                    for p in self.other:
                        para = document.add_paragraph(p["c"])
                        if p["h"]:
                            for run in para.runs:
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            self.l2tk_chart = []
            self.prospect_list = []
            self.other = []

        for (
            chart,
            position,
            artist,
            song,
            unsigned,
            l2tk,
            movement,
            days,
            pea,
            link,
            label,
            date,
        ) in final_df.itertuples(index=False):
            day = str(days).replace(".0", "")
            peak = str(pea).replace(".0", "")
            chart = chart.replace("- Freddy", "")
            if chart != chart_header:
                if chart_header:
                    add_content_and_header(chart_header, date)

                chart_header = chart

            if l2tk == "L2TK":
                # if (
                #     artist.lower() in prospect
                # ):  # Check if song (lowercased) is in the prospect list

                #     prospect_list.append(
                #         f"{position}. {artist}  -  {song} ({'=' if movement == '0' else movement})\n  • Days on chart: {day}\n  • Peak: {peak}\n"
                #     )
                # else:
                self.l2tk_chart.append(
                    f"{position}. {artist}  -  {song} ({'=' if movement == '0' else movement}) (L2TK)\n  • Days on chart: {day}\n  • Peak: {peak}\n"
                )

            if movement == "NEW" and unsigned == "UNSIGNED":
                self.other.append(
                    {
                        "c": f"{position}. {artist} -  {song} ({movement})\n  • Label: {label} (UNSIGNED)\n  • {link}\n",
                        "h": True,
                    }
                )

        add_content_and_header(chart_header, date)
        document.save(
            f'spotify_chart_intern_{datetime.datetime.now().strftime("%y%m%d")}.docx'
        )


def scrape_all():

    options = webdriver.ChromeOptions()
    # options.binary_location = "/opt/chrome/chrome"
    options.add_argument("--headless=new")
    # options.add_argument("--no-sandbox")
    # options.add_argument("--disable-gpu")
    # options.add_argument("--window-size=1963x1696")
    # options.add_argument("--single-process")
    # options.add_argument("--disable-dev-shm-usage")
    # options.add_argument("--disable-dev-tools")
    # options.add_argument("--no-zygote")
    # options.add_argument(f"--user-data-dir={mkdtemp()}")
    # options.add_argument(f"--data-path={mkdtemp()}")
    # options.add_argument(f"--disk-cache-dir={mkdtemp()}")
    # options.add_argument("--remote-debugging-port=9222")
    # service = webdriver.ChromeService("/opt/chromedriver")

    # local
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager

    service = Service(ChromeDriverManager().install())

    driver = webdriver.Chrome(service=service, options=options)
    scrape = Scrape(driver)
    scrape.spotify_signin()
    scrape.spotify(
        "SPOTIFY GLOBAL",
        "https://charts.spotify.com/charts/view/regional-global-daily/latest",
    )
    scrape.spotify(
        "SPOTIFY USA", "https://charts.spotify.com/charts/view/regional-us-daily/latest"
    )
    scrape.spotify(
        "SPOTIFY CA", "https://charts.spotify.com/charts/view/regional-ca-daily/latest"
    )
    scrape.spotify(
        "SPOTIFY UK",
        "https://charts.spotify.com/charts/view/regional-gb-daily/latest",
    )
    scrape.spotify(
        "SPOTIFY VIRAL GLOBAL",
        "https://charts.spotify.com/charts/view/viral-global-daily/latest",
    )
    scrape.spotify(
        "SPOTIFY VIRAL CA",
        "https://charts.spotify.com/charts/view/viral-CA-daily/latest",
    )
    scrape.spotify(
        "SPOTIFY VIRAL NZ",
        "https://charts.spotify.com/charts/view/viral-nz-daily/latest",
    )
    scrape.spotify(
        "SPOTIFY VIRAL UK",
        "https://charts.spotify.com/charts/view/viral-gb-daily/latest",
    )
    scrape.driver.quit()
    scrape.chart_search()
    scrape.create_docx()


def lambda_handler(event, context):
    scrape_all()
    return {
        "statusCode": 200,
        "body": "Scrape complete",
    }


lambda_handler(None, None)
